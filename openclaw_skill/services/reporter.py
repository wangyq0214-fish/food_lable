from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document

from openclaw_skill.models.schemas import AuditResult


class DocxReportService:
    def render(self, result: AuditResult, output_path: str) -> str:
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        output = self._make_writable_target(output)

        doc = Document()
        doc.add_heading("食品标签审核报告", level=1)
        doc.add_paragraph(f"文档ID: {result.document_id}")
        doc.add_paragraph(f"产品名称: {result.product_name or '未知'}")
        doc.add_paragraph(f"问题数量: {len(result.issues)}")

        doc.add_heading("问题清单", level=2)
        if not result.issues:
            doc.add_paragraph("未发现问题（当前为骨架流程结果）。")
        else:
            for idx, issue in enumerate(result.issues, start=1):
                doc.add_paragraph(
                    f"{idx}. [{issue.severity}] {issue.title} - {issue.description}"
                )
                if issue.suggestion:
                    doc.add_paragraph(f"   建议: {issue.suggestion}")

        doc.save(str(output))
        return str(output)

    def _make_writable_target(self, output: Path) -> Path:
        """Avoid PermissionError when target docx is locked (e.g. opened in Word)."""
        if not output.exists():
            return output

        try:
            with output.open("a", encoding="utf-8"):
                pass
            return output
        except OSError:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            return output.with_name(f"{output.stem}_{timestamp}{output.suffix}")
